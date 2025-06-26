from docx import Document

def write_summary(docs, technical_analyses, path="outputs/summary.docx"):
    """
    Write Word document with detailed technical analysis for procedures with complexity > 3.
    This document is intended to flag procedures that may need refactoring.
    """
    document = Document()
    
    # Add title and introduction
    document.add_heading("Stored Procedures Requiring Refactoring Review", level=1)
    intro = document.add_paragraph()
    intro.add_run("This document contains detailed technical analysis of stored procedures with complexity scores greater than 3. ")
    intro.add_run("These procedures have been flagged for potential refactoring to improve maintainability, performance, and code quality.")
    
    # Filter and add only high-complexity procedures
    high_complexity_count = 0
    for doc in docs:
        if doc["complexity"] > 3:
            high_complexity_count += 1
            
            # Find corresponding technical analysis
            tech_analysis = next((ta for ta in technical_analyses if ta["name"] == doc["sp_name"]), None)
            
            # Add procedure heading with complexity score
            document.add_heading(f'{doc["sp_name"]} (Complexity: {doc["complexity"]})', level=2)
            
            # Add business summary
            document.add_heading("Business Function", level=3)
            document.add_paragraph(doc["summary"])
            
            # Add technical analysis if available
            if tech_analysis:
                document.add_heading("Technical Analysis & Refactoring Recommendations", level=3)
                document.add_paragraph(tech_analysis["technical_analysis"])
            
            # Add complexity factors
            document.add_heading("Complexity Factors", level=3)
            factors_text = doc["complexity_factors"] if doc["complexity_factors"] else "No specific factors identified"
            document.add_paragraph(f"Lines of Code: {doc['lines_of_code']}")
            document.add_paragraph(f"Contributing Factors: {factors_text}")
            
            # Add separator
            document.add_paragraph("─" * 50)
    
    # Add summary at the end
    document.add_heading("Summary", level=2)
    summary_para = document.add_paragraph()
    summary_para.add_run(f"Total procedures analyzed: {len(docs)}")
    summary_para.add_run(f"\nProcedures flagged for refactoring review: {high_complexity_count}")
    summary_para.add_run(f"\nPercentage requiring attention: {(high_complexity_count/len(docs)*100):.1f}%" if len(docs) > 0 else "\nPercentage requiring attention: 0%")
    
    document.save(path)
    return high_complexity_count
